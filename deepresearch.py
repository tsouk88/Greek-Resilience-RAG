import os
import docx  
from langchain_google_genai import ChatGoogleGenerativeAI
from dotenv import load_dotenv

load_dotenv()

# 1. Σωστή ανάγνωση ολόκληρου του .docx αρχείου
def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    return "\n".join(full_text)

try:
    findings = read_docx("Κεφάλαιο 5 - Ανάλυση Ευαισθησίας, Μαθηματική Σύνθεση και Βιβλιογραφία RTIx V7 (Final).docx")
except Exception as e:
    print(f"Σφάλμα κατά την ανάγνωση του αρχείου: {e}")
    findings = ""

# 2. Αρχικοποίηση του Gemini 2.5 Pro με Live Google Search
llm = ChatGoogleGenerativeAI(
    model="gemini-2.5-pro",
    google_api_key=os.getenv("GEMMA_API_KEY"),
    temperature=1,
    thinking_budget=4096,
    model_kwargs={
        "tools": [{"google_search": {}}]
    }
)

prompt = f"""Είσαι κορυφαίος ακαδημαϊκός ερευνητής στην περιφερειακή οικονομική και την εξελικτική οικονομετρία.

Το καθήκον σου είναι να διεξάγεις μια ζωντανή αναζήτηση βιβλιογραφίας στο διαδίκτυο και να γράψεις μια προχωρημένη βιβλιογραφική ανασκόπηση επιπέδου διδακτορικού (PhD-level).

Το δείγμα των δεδομένων μου καλύπτει την περίοδο 2005–2022. Αυτό σημαίνει ότι το μοντέλο RTIx ΔΕΝ μετράει μόνο την πανδημία, αλλά αναλύει τη μακροχρόνια συμπεριφορά των ελληνικών περιφερειών απέναντι σε ΔΙΑΔΟΧΙΚΟΥΣ ΚΑΙ ΣΩΡΕΥΤΙΚΟΥΣ ΚΛΟΝΙΣΜΟΥΣ (Successive and Compounding Shocks): την Ελληνική Οικονομική Κρίση / Κρίση Χρέους της Ευρωζώνης (2008–2018) και την Πανδημία COVID-19 (2020–2022).

Χρησιμοποίησε το εργαλείο αναζήτησής σου (Google Search) για να βρεις πραγματικά, επαληθεύσιμα ακαδημαϊκά άρθρα (papers) που δημοσιεύτηκαν μεταξύ 2012 και 2026 σχετικά με:
1. Τη μακροχρόνια περιφερειακή οικονομική ανθεκτικότητα (long-term regional economic resilience) την κρίση χρεους του 2008.
2. Τη θεωρία των διαδοχικών ή σωρευτικών κλονισμών (compounding / consecutive shocks) στην Εξελικτική Οικονομική Γεωγραφία (Evolutionary Economic Geography).
3. Τη Θεωρία του Προσαρμοστικού Κύκλου (Adaptive Cycle / Panarchy) εφαρμοσμένη σε περιφέρειες που βιώνουν παρατεταμένη ύφεση και αναδιοργάνωση.
4. Την έννοια της Τροχιακής Απόκλισης (Trajectory Divergence) ή της δημιουργίας νέων μονοπατιών (new path creation) έναντι της εξάρτησης από την πορεία (path dependency) σε περιβάλλον πολλαπλών κρίσεων.

Στη συνέχεια, γράψε μια ολοκληρωμένη βιβλιογραφική ανασκόπηση (περίπου 500 λέξεις) ΣΤΑ ΕΛΛΗΝΙΚΑ, η οποία θα συνδέει άμεσα αυτές τις διεθνείς δημοσιεύσεις με τα δικά μου εμπειρικά ευρήματα (Rs, Rc, D) και τις συμπεριφορές των Clusters (π.χ. την Αττική ως Rigidity Trap με ακραία αρνητικό Rs λόγω της σωρευτικής κρίσης 2008-2022, και τις νησιωτικές περιφέρειες ως θύλακες Δημιουργικής Καταστροφής με υψηλό D).

Όλο το κείμενο της απάντησής σου πρέπει να είναι γραμμένο στην Ελληνική γλώσσα, με επίσημο, ακαδημαϊκό και αυστηρό επιστημονικό ύφος. Στο τέλος, παράθεσε τη βιβλιογραφία σε μορφή APA 7th Edition.

Εδώ είναι το πλήρες εμπειρικό κείμενο του Κεφαλαίου 5:
{findings}
"""

print("Εκτέλεση Deep Research και δημιουργία ελληνικής βιβλιογραφικής σύνθεσης...")
response = llm.invoke(prompt)

# 4. Αποθήκευση του αποτελέσματος σε Markdown
with open("literature_synthesis.md", "w", encoding="utf-8") as f:
    f.write(response.content)

print("Έτοιμο! Το αρχείο 'literature_synthesis.md' δημιουργήθηκε με επιτυχία στα ελληνικά.")